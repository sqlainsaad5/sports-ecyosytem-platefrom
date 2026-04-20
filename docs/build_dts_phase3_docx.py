"""
Generate UCP SDP Phase III — Design & Test Specification (DTS) for FYP26-CS-G22.
Output: DTS_Phase_III_Sports_Ecosystem_Platform_FILLED.docx
Run: python docs/build_dts_phase3_docx.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Design and Test Specification (DTS)")
    r.bold = True
    r.font.size = Pt(16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("SDP Phase III")
    r2.bold = True
    r2.font.size = Pt(14)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("Sports Ecosystem Platform").bold = True

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run("Faculty of Information Technology · Department of Computer Science · BSCS Final Project")

    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run("Group ID: ").bold = True
    t.add_run("FYP26-CS-G22")

    for label in (
        "Project Advisor: [Advisor Name]",
        "Student Name(s) / Reg. No(s).: [Fill as per Project Office record]",
        "Submission Date: April 2026",
    ):
        row = doc.add_paragraph()
        row.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.add_run(label)

    doc.add_page_break()


def add_revision_history(doc: Document) -> None:
    doc.add_heading("Revision History", level=1)
    table = doc.add_table(rows=4, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Version"
    hdr[1].text = "Date"
    hdr[2].text = "Author"
    hdr[3].text = "Reason For Changes"
    rows = [
        ("1.0", "Apr 2026", "FYP26-CS-G22", "Initial Phase III DTS aligned with implemented system"),
        ("", "", "", ""),
        ("", "", "", ""),
    ]
    for i, (v, d, a, r) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = v
        cells[1].text = d
        cells[2].text = a
        cells[3].text = r
    doc.add_paragraph()


def add_toc_placeholder(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph(
        "In Microsoft Word: References → Table of Contents → Automatic Table 1 "
        "(this file uses Heading 1–3 styles so the TOC can be generated or updated in one click)."
    )
    doc.add_page_break()


def add_body(doc: Document) -> None:
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This Design and Test Specification (DTS) for Phase III documents the completed and evolving "
        "design of the Sports Ecosystem Platform (FYP26-CS-G22), together with testing evidence and "
        "project status. It extends the Software Requirements Specification (Phase I) and Software "
        "Design Specification (Phase II) with implementation-level detail: architecture, components, "
        "interfaces, data architecture, external integrations, human–computer interaction choices, "
        "workflows, test planning and results, and defect tracking."
    )
    doc.add_paragraph(
        "The product is a centralized web application (not a research dataset, game executable, or "
        "hardware prototype). It applies established algorithms for matching and scheduling in the "
        "application domain of sports management, and integrates third-party payment services where configured."
    )

    doc.add_heading("1.1 Product overview and problem statement", level=2)
    doc.add_paragraph(
        "Stakeholders—players, coaches, equipment vendors, and administrators—typically rely on "
        "informal networks, phone calls, and disconnected tools for coaching discovery, facility booking, "
        "training records, and retail. The Sports Ecosystem Platform consolidates these concerns into one "
        "role-based system: automated coach recommendation, indoor ground booking for cricket and badminton, "
        "training session lifecycle (requests, schedules, attendance, weekly performance points, training plans), "
        "a subscription-quota equipment marketplace, payments, verification and complaints workflows, "
        "notifications, and administrative reporting."
    )

    doc.add_heading("1.2 Objectives", level=2)
    objectives = [
        "Deliver a secure multi-role web platform with JWT authentication and role-based access control.",
        "Implement automated coach recommendations using player profile attributes (sport, skill, location) and coach availability.",
        "Provide real-time indoor ground booking with short payment holds to prevent double booking.",
        "Support training management: requests, sessions, weekly plans, attendance, and performance evaluation.",
        "Enable business owners to sell sports equipment with listing quotas tied to subscription packages.",
        "Integrate card payments via Stripe PaymentIntents when keys are configured; retain development fallbacks without keys.",
        "Provide administrator verification, monitoring, complaints handling, and summary reporting.",
    ]
    for o in objectives:
        doc.add_paragraph(o, style="List Bullet")

    doc.add_heading("1.3 Scope of the product", level=2)
    doc.add_paragraph(
        "In scope: browser-based SPA and REST API; MongoDB persistence; file uploads for documents and "
        "product images; in-app notifications; admin configuration of sports and grounds; Stripe-backed "
        "payments when enabled."
    )
    doc.add_paragraph(
        "Out of scope (per SRS/SDD): native mobile apps, live video coaching, third-party logistics "
        "integration, full social network features, and sports beyond the configured categories "
        "(initially cricket and badminton as seeded categories)."
    )

    doc.add_heading("2. Abstract", level=1)
    doc.add_paragraph(
        "The Sports Ecosystem Platform addresses fragmentation in local sports activity by connecting "
        "players, coaches, business owners, and administrators in one ecosystem. The significance lies in "
        "transparent scheduling, traceable performance history, and trustworthy verification of coaches "
        "and vendors. Knowledge areas include full-stack web engineering (React, Node.js, Express, MongoDB), "
        "security (hashing, JWT, RBAC), REST API design, payment integration, and usability for non-technical users."
    )
    doc.add_paragraph(
        "The expected result is a deployable system with documented APIs, automated integration tests for "
        "core auth flows, and clear operational guidance (environment variables, database seeding, CORS, rate limits)."
    )

    doc.add_heading("3. Document conventions", level=1)
    doc.add_paragraph(
        "Headings use Word built-in styles (Heading 1–3). Acronyms are expanded on first use in each major "
        "section where helpful. API paths are written as /api/...; use cases from the SRS are referenced as "
        "UC-P*, UC-C*, UC-B*, UC-A* where applicable. Code and collection names appear in monospace in this "
        "document when created from the generator; you may switch to Consolas for code in Word if desired."
    )

    doc.add_heading("4. Technical architecture", level=1)
    doc.add_paragraph(
        "The system follows a layered client–server architecture. The presentation tier is a React 19 single-page "
        "application built with Vite 8 and styled with Tailwind CSS 3. It communicates over HTTPS (dev: Vite with "
        "optional basic SSL plugin; production: HTTPS at reverse proxy) with a JSON REST API."
    )
    doc.add_paragraph(
        "The application tier is Node.js 18+ with Express 4. Cross-cutting concerns include CORS, JSON body "
        "parsing (2 MB limit), structured logging via morgan, global error handling, static serving of uploads, "
        "and express-rate-limit on /api. Authentication uses bcrypt-hashed passwords and JWT bearer tokens."
    )
    doc.add_paragraph(
        "The data tier is MongoDB accessed through Mongoose 8 models. Binary artifacts (documents, images) "
        "are stored on the server filesystem under backend/uploads and exposed read-only at /uploads."
    )

    doc.add_heading("4.1 High-level architecture (major subsystems)", level=2)
    doc.add_paragraph("Major subsystems:")
    for item in [
        "Web client (React Router, role-specific layouts, Axios API service with JWT).",
        "API gateway (Express Router under /api with rate limiting).",
        "Domain modules: auth, player, coach, business, admin, public catalog, uploads.",
        "Payment adapter (Stripe PaymentIntents when STRIPE_SECRET_KEY is set; otherwise mock/last-four style flows for local dev).",
        "Notification subsystem (in-app records; email SMTP hooks described for extension).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "Figure (to paste from draw.io / Visio): a diagram with boxes “Browser SPA”, “REST API”, "
        "“MongoDB”, “File store”, “Stripe”, and arrows labeled HTTPS, BSON queries, static files, and payment API."
    )

    doc.add_heading("5. Technology architecture and infrastructure", level=1)
    doc.add_paragraph(
        "Development workstations run Windows (or similar) with Node.js 18+, npm, and optionally MongoDB "
        "Community locally or MongoDB Atlas in the cloud. The API default port is 5000; the Vite dev server "
        "uses 5173 and proxies /api and /uploads to the backend (see frontend/vite.config.js)."
    )
    doc.add_paragraph(
        "Production expectations: host the API behind HTTPS (nginx, Caddy, or cloud load balancer), set "
        "CLIENT_URL for CORS, use strong JWT_SECRET, enable MongoDB authentication and backups, and configure "
        "Stripe live keys only in production secrets."
    )
    doc.add_paragraph(
        "Modes of operation: online-only; users access the SPA in a modern browser (Chrome, Edge, Firefox). "
        "No thick client or offline-first sync is required for the MVP."
    )

    doc.add_heading("5.1 Rationale for technology choices", level=2)
    doc.add_paragraph(
        "React with Vite was chosen for fast development and a rich component ecosystem; Tailwind for "
        "consistent responsive UI. Node.js and Express align with JSON REST services and share JavaScript "
        "with the frontend for team skill efficiency. MongoDB provides flexible document modeling for "
        "profiles, bookings, and orders. Stripe is a widely used PCI-aware card processing provider with "
        "a maintained Node SDK and frontend loaders (@stripe/stripe-js, @stripe/react-stripe-js)."
    )

    doc.add_heading("5.2 Architecture evaluation", level=2)
    doc.add_paragraph(
        "Strengths: clear separation between SPA and API; JWT statelessness simplifies horizontal scaling "
        "of API instances; Mongoose schemas document data shape; Stripe reduces PCI scope for card data."
    )
    doc.add_paragraph(
        "Trade-offs: MongoDB requires careful indexing for booking overlap queries; file storage on local disk "
        "should move to object storage for multi-instance deployment; email notifications need SMTP integration "
        "for production alerts."
    )

    doc.add_heading("6. Application and data architecture", level=1)
    doc.add_paragraph(
        "Application components map to REST resources under /api/auth, /api/players, /api/coaches, "
        "/api/business, /api/admin, /api/public, and /api/uploads. Data components are MongoDB collections "
        "including users, role profiles, grounds, bookings, training entities, commerce entities, "
        "notifications, complaints, verification documents, and settings (see repository docs/SETUP_AND_DATABASE.md)."
    )

    doc.add_heading("6.1 Logical data model (summary)", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    h = t.rows[0].cells
    h[0].text = "Collection / area"
    h[1].text = "Role / purpose"
    rows = [
        ("users", "Authentication, role, verification and suspension flags"),
        ("playerprofiles, coachprofiles, businessprofiles", "Role-specific attributes and links"),
        ("sportcategories, indoorgrounds", "Admin-managed taxonomy and facilities"),
        ("groundbookings", "Hold/confirm/cancel with hold expiry"),
        ("trainingrequests, trainingsessions, trainingplans", "Training lifecycle"),
        ("attendancerecords, performanceevaluations", "Session attendance and weekly scores"),
        ("products, orders, payments", "E-commerce and typed payments"),
        ("notifications, complaints, verificationdocuments", "Cross-cutting workflows"),
        ("coachfeedbacks, coachpartnershiprequests", "Feedback and B2B requests"),
        ("systemsettings", "Configurable platform parameters"),
    ]
    for a, b in rows:
        row = t.add_row().cells
        row[0].text = a
        row[1].text = b

    doc.add_heading("7. Component design", level=1)
    doc.add_paragraph(
        "Controllers orchestrate HTTP I/O; Mongoose models encapsulate validation and persistence; middleware "
        "handles authenticate, requireRole, validate (express-validator), and multer uploads. This mirrors a "
        "lightweight service layer pattern without unnecessary abstraction for the FYP scope."
    )

    doc.add_heading("7.1 Design patterns and reuse", level=2)
    for p in [
        "Layered architecture (routes → middleware → controllers → models).",
        "RBAC middleware for every protected route.",
        "Gateway-style single /api entry with rate limiting.",
        "Quota counter pattern for business listing limits per subscription tier.",
        "Hold-expiry pattern for ground booking payment windows.",
    ]:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("8. Component interactions and interfaces", level=1)

    doc.add_heading("8.1 Component–component interfaces", level=2)
    doc.add_paragraph(
        "Internal interfaces are HTTP + JSON between SPA and API, and Mongoose calls between controllers "
        "and MongoDB. Sequence flows: login returns JWT; subsequent requests send Authorization: Bearer <token>."
    )

    doc.add_heading("8.2 Component–external system interfaces", level=2)
    doc.add_paragraph(
        "Stripe: when STRIPE_SECRET_KEY is configured, the backend creates or verifies PaymentIntents; the "
        "frontend collects card details through Stripe.js elements where integrated (see StripePaySection.jsx "
        "and controllers using utils/stripePayments.js). No NADRA or bank core integration is implemented; "
        "card processing is delegated to Stripe."
    )
    doc.add_paragraph(
        "MongoDB Atlas (optional): standard SRV connection string via MONGODB_URI."
    )

    doc.add_heading("8.3 Component–human interface", level=2)
    doc.add_paragraph(
        "Humans interact through the React UI: landing and auth pages; player/coach/business/admin dashboards "
        "with side navigation; forms and tables for CRUD and workflows. JWT persistence is in browser storage; "
        "session expiry redirects to login via Axios interceptors."
    )

    doc.add_heading("9. Related and previous work", level=1)
    doc.add_paragraph(
        "Commercial sports-tech products (generic examples: marketplace + booking platforms) combine commerce "
        "and scheduling but often omit FYP-specific coach verification and academic traceability. Academic "
        "SRS/SDD documents for this project (Phase I–II) define use cases UC-P/C/B/A that this implementation "
        "maps to REST endpoints in docs/API_ENDPOINTS.md. The implementation prioritizes clarity and "
        "demonstrability over horizontal scale."
    )

    doc.add_heading("10. Human–computer interaction (HCI)", level=1)
    doc.add_paragraph(
        "The UI uses Tailwind for spacing, typography, and responsive grids. Each role has a dedicated layout "
        "to reduce cognitive load. Primary tasks (book ground, request training, manage products) are reachable "
        "within a few clicks from each dashboard. Error messages from the API are surfaced in context where "
        "implemented."
    )

    doc.add_heading("11. Detailed design and workflows", level=1)
    doc.add_paragraph(
        "Detailed API behavior is specified in docs/API_ENDPOINTS.md. Key workflows:"
    )
    workflows = [
        "Player registration → profile → coach recommendations → training request → coach accept → sessions and plans.",
        "Ground booking: list availability → hold slot (~5 min) → pay → confirm booking; hold expires if unpaid.",
        "Business: subscribe/renew → listing quota → product CRUD → orders and status updates.",
        "Admin: verify documents, manage users/grounds/sports, monitor bookings and performance, resolve complaints.",
    ]
    for w in workflows:
        doc.add_paragraph(w, style="List Bullet")

    doc.add_paragraph(
        "Swim-lane diagram: draw with lanes Player / Coach / Business / Admin / System showing the flows above."
    )
    doc.add_paragraph(
        "Design-level sequence diagrams (auth, booking hold/confirm, subscription purchase) should be drawn "
        "in Phase III with more detail than Phase II SDS; attach as figures or appendix pages."
    )

    doc.add_heading("11.1 Other design details", level=2)
    doc.add_paragraph(
        "Research-oriented deliverables: this project is an engineering product; experimental results and "
        "statistical analysis are not primary outputs. Hardware dependencies: standard PCs, network access, "
        "and browser—no custom embedded devices. Game engine or real-time multiplayer servers are out of scope."
    )

    doc.add_heading("11.2 Feedback from previous phases (SRS / SDS)", level=2)
    doc.add_paragraph(
        "Record supervisor or evaluation committee feedback from Phase I–II defenses and how this version "
        "of the system and documentation addresses it."
    )
    ft = doc.add_table(rows=1, cols=3)
    ft.style = "Table Grid"
    fc = ft.rows[0].cells
    fc[0].text = "Comment / feedback (Phase I–II)"
    fc[1].text = "How addressed in Phase III"
    fc[2].text = "Status"
    for _ in range(3):
        row = ft.add_row().cells
        row[0].text = "[Fill]"
        row[1].text = "[Fill]"
        row[2].text = "[Fill]"

    doc.add_heading("12. Test specification and results", level=1)
    doc.add_paragraph(
        "Automated integration tests (Jest + Supertest) reside in backend/tests/api.integration.test.js "
        "and require MONGODB_URI in backend/.env."
    )

    doc.add_heading("12.1 Sample test cases", level=2)
    tt = doc.add_table(rows=1, cols=6)
    tt.style = "Table Grid"
    th = tt.rows[0].cells
    headers = ["ID", "Requirement / UC", "Pre-condition", "Steps", "Expected", "Result"]
    for i, h in enumerate(headers):
        th[i].text = h
    cases = [
        (
            "TC-01",
            "API health",
            "Server running",
            "GET /api/health",
            "200, success true",
            "Pass (automated)",
        ),
        (
            "TC-02",
            "UC-P1 register",
            "Unique email",
            "POST /api/auth/register player payload",
            "201 + JWT",
            "Pass (automated)",
        ),
        (
            "TC-03",
            "Validation",
            "Coach without specialties",
            "POST /api/auth/register coach empty specialties",
            "400 validation error",
            "Pass (automated)",
        ),
        (
            "TC-04",
            "Auth",
            "Registered user",
            "POST /api/auth/login wrong password",
            "401",
            "Pass (automated)",
        ),
        (
            "TC-05",
            "UC-P6 ground",
            "Player logged in, ground data",
            "Hold → pay → confirm (manual / Stripe test)",
            "Booking confirmed, no overlap",
            "Pass manual / env-dependent",
        ),
        (
            "TC-06",
            "UC-B6 product quota",
            "Business at quota",
            "POST product",
            "Rejected or message per business rules",
            "Verify manual",
        ),
    ]
    for row_data in cases:
        row = tt.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph(
        "Summary of test results: automated tests cover health, registration validation, successful player signup, "
        "and failed login. Broader user acceptance testing should follow the SRS acceptance scenarios per role "
        "(player, coach, business owner, admin)."
    )

    doc.add_heading("13. Defect tracking", level=1)
    dt = doc.add_table(rows=1, cols=4)
    dt.style = "Table Grid"
    dh = dt.rows[0].cells
    for i, h in enumerate(["Module", "Defect description", "Severity", "Status"]):
        dh[i].text = h
    example = (
        "—",
        "Cosmetic / minor issues tracked via repository issues during polish (fill as found)",
        "Low",
        "Open / closed as applicable",
    )
    r = dt.add_row().cells
    for i, v in enumerate(example):
        r[i].text = v

    doc.add_heading("14. Project completion status", level=1)
    doc.add_paragraph(
        "Insert a Gantt chart developed in Microsoft Project or Excel per Project Office guidelines. "
        "Below is a summary table of modules versus implementation status in this repository snapshot."
    )
    mt = doc.add_table(rows=1, cols=3)
    mt.style = "Table Grid"
    mh = mt.rows[0].cells
    mh[0].text = "Module / area"
    mh[1].text = "Status"
    mh[2].text = "Notes"
    modules = [
        ("Authentication & RBAC", "Complete", "JWT, bcrypt, role routes"),
        ("Player features (coaches, training, grounds, shop, orders, performance, complaints)", "Complete", "See /player routes"),
        ("Coach features (requests, sessions, plans, performance, feedback, payments, docs)", "Complete", "See /coach routes"),
        ("Business features (products, orders, subscription, coaches, docs)", "Complete", "Quota + Stripe when configured"),
        ("Admin features (verification, users, sports, grounds, monitoring, reports)", "Complete", "AdminLayout routes"),
        ("Payments (Stripe)", "Complete when keys set", "Fallback without keys for dev"),
        ("Email notifications", "Partial / stub", "Extend notify utility for SMTP"),
        ("Automated test suite", "Partial", "Integration tests; expand for full regression"),
    ]
    for a, b, c in modules:
        row = mt.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c

    doc.add_heading("15. Screenshots", level=1)
    doc.add_paragraph(
        "Include full-window screenshots for representative screens, labeled by role and feature. Suggested list:"
    )
    shots = [
        "Landing, Login, Register",
        "Player: dashboard, coaches, training, grounds, shop, orders, performance, notifications, complaint, profile",
        "Coach: dashboard, requests, sessions, plans, performance, matches (recommendations), feedback, payments, documents, notifications",
        "Business: dashboard, products, orders, subscription, coaches, documents, notifications",
        "Admin: dashboard, verification (coaches/business), directory, users, sports, grounds, monitor bookings/performance, reports, subscriptions, complaints, settings",
    ]
    for s in shots:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("16. References", level=1)
    refs = [
        "React documentation — https://react.dev/",
        "Vite documentation — https://vite.dev/",
        "Express.js guide — https://expressjs.com/",
        "MongoDB Manual — https://www.mongodb.com/docs/manual/",
        "Mongoose documentation — https://mongoosejs.com/docs/guide.html",
        "Stripe API reference — https://stripe.com/docs/api",
        "Project internal: docs/API_ENDPOINTS.md, docs/PHASE2_SYSTEM_DESIGN.md, docs/SETUP_AND_DATABASE.md, README.md",
        "University of Central Punjab — FYP26-CS-G22 SRS and SDD (Phase I–II)",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("Appendix A — Glossary", level=1)
    gloss = [
        ("API", "Application Programming Interface; here, REST under /api."),
        ("JWT", "JSON Web Token used for bearer authentication."),
        ("RBAC", "Role-based access control (player, coach, business_owner, admin)."),
        ("SPA", "Single-page application."),
        ("UC", "Use case identifier from SRS."),
        ("MVP", "Minimum viable product."),
        ("SRS", "Software Requirements Specification (Phase I)."),
        ("SDS/SDD", "Software Design Specification (Phase II)."),
        ("DTS", "Design and Test Specification (Phase III)."),
    ]
    gt = doc.add_table(rows=1, cols=2)
    gt.style = "Table Grid"
    gh = gt.rows[0].cells
    gh[0].text = "Term"
    gh[1].text = "Definition"
    for term, defin in gloss:
        row = gt.add_row().cells
        row[0].text = term
        row[1].text = defin

    doc.add_heading("Appendix B — Additional information", level=1)
    doc.add_paragraph(
        "Repository layout: backend/ (Express API), frontend/ (Vite React SPA), docs/ (analysis and API reference). "
        "Scripts: npm run dev from root runs API and web concurrently; npm test runs backend tests."
    )
    doc.add_paragraph(
        "This generated document is the Phase III DTS content aligned with the Sports Ecosystem Platform. "
        "Personal names on the cover sheet must match Project Office records. Figures (architecture, swim-lane, Gantt) "
        "should be pasted from your diagram tools to satisfy visual submission requirements."
    )


def main() -> None:
    doc = Document()
    set_normal_style(doc)
    add_title_block(doc)
    add_revision_history(doc)
    add_toc_placeholder(doc)
    add_body(doc)

    out = Path(__file__).resolve().parent / "DTS_Phase_III_Sports_Ecosystem_Platform_FILLED.docx"
    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
